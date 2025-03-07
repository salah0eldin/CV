import json
import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx2pdf import convert

# For the horizontal line hack:
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

specialities = ["software", "embedded", "digital", "web"]

#
# ------------- STYLING HELPERS -------------
#


def set_base_style(document):
    """
    Sets the base 'Normal' style:
      - Arial
      - 11 pt
      - Single spacing
      - 0 spacing before/after paragraphs
    """
    style = document.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)


def set_page_margins(document, margin_in_inches=0.2):
    """
    Sets uniform page margins (top, bottom, left, right).
    Default is 0.3 inches for a compact look.
    """
    for section in document.sections:
        section.top_margin = Inches(0.1)
        section.bottom_margin = Inches(margin_in_inches)
        section.left_margin = Inches(margin_in_inches)
        section.right_margin = Inches(margin_in_inches)


def add_horizontal_line(document):
    """
    Inserts a horizontal line (by applying a bottom border to an empty paragraph).
    This is a python-docx XML trick.
    """
    p = document.add_paragraph()
    p.style.font.size = Pt(3)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "5")  # Thickness
    bottom.set(qn("w:space"), "0")  # Spacing
    bottom.set(qn("w:color"), "000000")

    pBdr.append(bottom)
    pPr.append(pBdr)


#
# ------------- CONTENT HELPERS -------------
#


def add_name_header(document, name_text):
    """
    Creates a centered paragraph for the name:
      - Uppercase
      - Bold
      - ~24pt
    """
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name_text.upper())
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(24)


def add_centered_line(document, text, font_size=11, bold=False):
    """
    Adds a single centered line (for contact info, links, etc.).
    """
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(font_size)


def add_section_heading(document, heading_text):
    """
    Adds a left-aligned heading in uppercase, bold, ~13pt.
    Followed by one blank line.
    """
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(heading_text.upper())
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(13)
    # Add a single blank line after the heading
    # document.add_paragraph("")


def add_bullet_paragraph(
    document,
    label="",
    text="",
    bullet_symbol="•",
    indent_left_inches=0.4,
    indent_right_inches=0,
    indent_hang_inches=0,
    bold_label=True,
    justify=True,
    bold=False,
    space_after=0,
):
    """
    Adds a sub-bullet paragraph with a round bullet (default '•' U+2022), "▪"
    often used for 'Description' or 'Key Elements'.
    'label' (e.g. 'Description: ') can be bold, while 'text' is normal.
    """
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent_left_inches)
    p.paragraph_format.right_indent = Inches(indent_right_inches)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(-indent_hang_inches)
    p.paragraph_format.space_after = Pt(space_after)

    # Bullet run
    if bullet_symbol != "":
        bullet_run = p.add_run(f"{bullet_symbol} ")
        bullet_run.font.name = "Arial"
        bullet_run.font.size = Pt(11)

    if label != "":
        # Label run (bold if needed)
        label_run = p.add_run(label)
        if bold_label:
            label_run.bold = True
        label_run.font.name = "Arial"
        label_run.font.size = Pt(11)

    # Text run (normal)
    text_run = p.add_run(text)
    text_run.font.name = "Arial"
    text_run.font.size = Pt(11)
    if bold:
        text_run.bold = True


def add_skills_table(document, skills, specialty_order):
    """
    Adds a table for the skills section without a header row.
    Each row contains:
      - The specialty (capitalized, bold, 11pt, right aligned) in the first column.
      - Its skills (joined with " - ") in the second column, in 11pt.
    The table is centered on the page.
    All borders are removed.
    """
    # Create a table with no initial rows and 2 columns.
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center the whole table

    for sp in specialty_order:
        if sp in skills:
            row = table.add_row()
            cell0 = row.cells[0]
            cell1 = row.cells[1]

            # First cell: specialty name, bold, 11pt, right aligned.
            p0 = cell0.paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p0.paragraph_format.left_indent = Inches(0.2)
            run0 = p0.add_run("▪ " + sp.capitalize() + ":")
            run0.bold = True
            run0.font.name = "Arial"
            run0.font.size = Pt(11)
            # Second cell: skills list, 11pt, left aligned.
            p1 = cell1.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p1.paragraph_format.space_after = Pt(3)
            run1 = p1.add_run(" - ".join(skills[sp]))
            run1.font.name = "Arial"
            run1.font.size = Pt(11)

            # Adjust column widths (first column remains narrow, second column is increased)
            cell0.width = Inches(1.3)
            cell1.width = Inches(6.8)

    # Remove all borders.
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = tblBorders.find(qn("w:" + border_name))
        if border is None:
            border = OxmlElement("w:" + border_name)
            tblBorders.append(border)
        border.set(qn("w:val"), "nil")

    return table


#
# ------------- MAIN CV BUILDER -------------
#


def build_cv(document, data, specialty=None):
    """
    Builds either a general CV (if specialty is None) or a specialized CV.
    Matches the layout from your screenshot:
      - Name (centered, uppercase, bold, 24pt)
      - Contact info (centered)
      - Horizontal line
      - SUMMARY, EDUCATION, WORK EXPERIENCE, SKILLS (as a table), PROJECTS, COURSES
    Uses:
      - Smaller square bullets (▪) for top-level items
      - Round bullets (•) for sub-items (Description, Key Elements)
      - Minimal spacing
    """

    # ---------- HEADER ----------
    add_name_header(document, data["header"]["name"])
    add_centered_line(document, data["header"]["contact"], font_size=10, bold=True)
    add_centered_line(document, data["header"]["links"], font_size=10, bold=True)
    add_horizontal_line(document)

    # ---------- SUMMARY ----------
    summary_list = data.get("summary", [])
    if summary_list:
        add_section_heading(document, "Summary")
        for summary_item in summary_list:
            add_bullet_paragraph(
                document,
                label="",
                text=summary_item,
                bullet_symbol="",
                indent_left_inches=0.2,
                indent_right_inches=0.2,
            )
    add_horizontal_line(document)

    # ---------- EDUCATION ----------
    education_list = data.get("education", [])
    if education_list:
        add_section_heading(document, "Education")
        for edu_item in education_list:
            add_bullet_paragraph(
                document,
                label="",
                text=edu_item,
                bullet_symbol="▪",
                indent_left_inches=0.2,
                indent_right_inches=0.2,
                bold=True,
            )
    add_horizontal_line(document)

    # ---------- WORK EXPERIENCE ----------
    work_experience_list = data.get("work_experience", [])
    if work_experience_list:
        add_section_heading(document, "Work Experience")
        for work_item in work_experience_list:
            add_bullet_paragraph(
                document,
                label="",
                text=work_item["header"],
                bullet_symbol="▪",
                indent_left_inches=0.2,
                justify=False,
                bold=True,
            )
            add_bullet_paragraph(
                document,
                "",
                work_item["body"],
                bullet_symbol="•",
                indent_left_inches=0.4,
                indent_right_inches=0.2,
            )
    add_horizontal_line(document)

    # ---------- SKILLS (as a table) ----------
    add_section_heading(document, "Skills")
    if specialty:
        specialty_order = [specialty] + [s for s in specialities if s != specialty]
    else:
        specialty_order = specialities
    add_skills_table(document, data["skills"], specialty_order)
    add_horizontal_line(document)

    # ---------- PROJECTS ----------
    add_section_heading(document, "Projects")
    projects = data.get("projects", {})
    if specialty:
        if specialty in projects:
            for proj in projects[specialty].get("main", []):
                add_bullet_paragraph(
                    document, proj["title"], bullet_symbol="▪", indent_left_inches=0.3
                )
                add_bullet_paragraph(
                    document,
                    "Description: ",
                    proj["detailed"],
                    bullet_symbol="•",
                    indent_left_inches=0.6,
                )
                add_bullet_paragraph(
                    document,
                    "Key Elements: ",
                    proj["key_elements"],
                    bullet_symbol="•",
                    indent_left_inches=0.6,
                )
        others = [s for s in specialities if s != specialty]
        for sp_other in others:
            if sp_other in projects:
                for proj in projects[sp_other].get("shortend", []):
                    add_bullet_paragraph(
                        document,
                        proj["title"],
                        bullet_symbol="▪",
                        indent_left_inches=0.3,
                    )
                    add_bullet_paragraph(
                        document,
                        "Description: ",
                        proj["detailed"],
                        bullet_symbol="•",
                        indent_left_inches=0.6,
                    )
                    add_bullet_paragraph(
                        document,
                        "Key Elements: ",
                        proj["key_elements"],
                        bullet_symbol="•",
                        indent_left_inches=0.6,
                    )
    else:
        for sp_general in specialities:
            if sp_general in projects:
                for proj in projects[sp_general].get("main", []):
                    add_bullet_paragraph(
                        document,
                        label="",
                        text=proj["title"],
                        bullet_symbol="▪",
                        indent_left_inches=0.3,
                    )
                    add_bullet_paragraph(
                        document,
                        "Description: ",
                        proj["detailed"],
                        bullet_symbol="•",
                        indent_left_inches=0.6,
                    )
                    add_bullet_paragraph(
                        document,
                        "Key Elements: ",
                        proj["key_elements"],
                        bullet_symbol="•",
                        indent_left_inches=0.6,
                    )
    add_horizontal_line(document)

    # ---------- COURSES ----------
    add_section_heading(document, "Courses")
    courses = data.get("courses", {})
    if specialty:
        if specialty in courses:
            for course in courses[specialty].get("main", []):
                add_bullet_paragraph(
                    document, course["title"], bullet_symbol="▪", indent_left_inches=0.3
                )
                add_bullet_paragraph(
                    document,
                    "Description: ",
                    course["detailed"],
                    bullet_symbol="•",
                    indent_left_inches=0.6,
                )
        others = [s for s in specialities if s != specialty]
        for sp_other in others:
            if sp_other in courses:
                for course in courses[sp_other].get("shortend", []):
                    add_bullet_paragraph(
                        document,
                        course["title"],
                        bullet_symbol="▪",
                        indent_left_inches=0.3,
                    )
                    add_bullet_paragraph(
                        document,
                        "Description: ",
                        course["detailed"],
                        bullet_symbol="•",
                        indent_left_inches=0.6,
                    )
    else:
        for sp_general in specialities:
            if sp_general in courses:
                for course in courses[sp_general].get("main", []):
                    add_bullet_paragraph(
                        document,
                        course["title"],
                        bullet_symbol="▪",
                        indent_left_inches=0.3,
                    )
                    add_bullet_paragraph(
                        document,
                        "Description: ",
                        course["detailed"],
                        bullet_symbol="•",
                        indent_left_inches=0.6,
                    )
    add_horizontal_line(document)


#
# ------------- MAIN FUNCTIONS -------------
#


def generate_cv_files(data):
    """
    Generates:
      1) General CV => 'Salah_CV.docx' (+ PDF)
      2) Specialized CVs => 'Salah_Software_CV.docx' (+ PDF), etc.
    """
    # General CV
    doc = Document()
    set_page_margins(doc, margin_in_inches=0.3)
    set_base_style(doc)
    build_cv(doc, data, specialty=None)
    general_filename = "Salah_CV.docx"
    doc.save(general_filename)
    print(f"Saved {general_filename}")
    try:
        convert(general_filename)
        print(f"Converted {general_filename} to PDF")
    except Exception as e:
        print("PDF conversion failed:", e)

    # Specialized CVs
    # for sp in specialities:
    for sp in ["software"]:
        doc = Document()
        set_page_margins(doc, margin_in_inches=0.3)
        set_base_style(doc)
        build_cv(doc, data, specialty=sp)
        filename = f"Salah_{sp.capitalize()}_CV.docx"
        doc.save(filename)
        print(f"Saved {filename}")
        try:
            convert(filename)
            print(f"Converted {filename} to PDF")
        except Exception as e:
            print("PDF conversion failed:", e)


#
# ------------------- ENTRY POINT -------------------
#

if __name__ == "__main__":
    data_file = "./data.json"
    with open(data_file, "r", encoding="utf-8") as f:
        data = json.load(f)

    specialities = data.get("specialities", specialities)

    generate_cv_files(data)
